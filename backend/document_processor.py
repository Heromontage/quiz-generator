import PyPDF2
import docx
from docx import Document
import io
import re

class DocumentProcessor:
    """Process various document formats and extract text"""
    
    SUPPORTED_FORMATS = ["pdf", "docx", "txt", "doc"]
    
    def extract_text(self, content: bytes, filename: str) -> str:
        """
        Extract text from various document formats
        
        Args:
            content: File content as bytes
            filename: Filename to determine format
        
        Returns:
            Extracted text
        """
        file_extension = filename.split('.')[-1].lower()
        
        try:
            if file_extension == "pdf":
                return self._extract_from_pdf(content)
            elif file_extension == "docx":
                return self._extract_from_docx(content)
            elif file_extension in ["txt"]:
                return content.decode('utf-8', errors='ignore')
            elif file_extension == "doc":
                return self._extract_from_doc(content)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
        except Exception as e:
            raise Exception(f"Error processing {filename}: {str(e)}")
    
    def _extract_from_pdf(self, content: bytes) -> str:
        """Extract text from PDF"""
        try:
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text()
            return text
        except Exception as e:
            raise Exception(f"Error reading PDF: {str(e)}")
    
    def _extract_from_docx(self, content: bytes) -> str:
        """Extract text from DOCX"""
        try:
            doc_file = io.BytesIO(content)
            doc = Document(doc_file)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text
        except Exception as e:
            raise Exception(f"Error reading DOCX: {str(e)}")
    
    def _extract_from_doc(self, content: bytes) -> str:
        """
        Extract text from DOC (using python-docx for legacy format)
        Note: For better DOC support, consider using python-pptx or other libraries
        """
        try:
            doc_file = io.BytesIO(content)
            doc = Document(doc_file)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text
        except Exception as e:
            raise Exception(f"Error reading DOC: {str(e)}")
    
    def clean_text(self, text: str) -> str:
        """Clean extracted text"""
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove special characters but keep punctuation
        text = re.sub(r'[^\w\s\.\,\!\?\-\:\;]', '', text)
        return text.strip()